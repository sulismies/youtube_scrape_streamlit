# Get channel ids
# https://commentpicker.com/youtube-channel-id.php
#
# quotas
# https://console.cloud.google.com/apis/api/youtube.googleapis.com/quotas?hl=en&project=youtube-scrape-dota2&pageState=(%22allQuotasTable%22:(%22f%22:%22%255B%255D%22,%22s%22:%5B(%22i%22:%22displayDimensions%22,%22s%22:%220%22),(%22i%22:%22currentPercent%22,%22s%22:%221%22),(%22i%22:%22sevenDayPeakPercent%22,%22s%22:%220%22),(%22i%22:%22currentUsage%22,%22s%22:%221%22),(%22i%22:%22sevenDayPeakUsage%22,%22s%22:%220%22),(%22i%22:%22serviceTitle%22,%22s%22:%220%22),(%22i%22:%22displayName%22,%22s%22:%220%22)%5D))
#
# ChatGPT Thread
# https://chatgpt.com/c/6720af1e-df2c-800f-a351-17239beb5ba2?model=o1-mini
#
#
# Youtube api 
# AIzaSyC2MmBms11
# CcqAeFhLXrtWa2K
# python scrape.py UCZsM8MOy0VC9blj_wBkbo-g 5 2024-10-05 
# python scrape.py UCwI9DhoGEziLUxTpK8H77jw 
# python scrape.py UCvTcxoyItMUSlw8T2MajftA 100 2023-10-30
# python scrape.py UCUqLL4VcEy4mXcQL0O_H_bg 1 2023-10-30
# python scrape.py UCy0-ftAwxMHzZc74OhYT3PA 1 2023-10-30
# python scrape.py UCUqLL4VcEy4mXcQL0O_H_bg 1 2023-10-30
# PvFKU8vDU

# channel_identifiers = [
#     'UCZsM8MOy0VC9blj_wBkbo-g',   # https://www.youtube.com/@PurgeGamers               OK
#     "UCwI9DhoGEziLUxTpK8H77jw",   # https://www.youtube.com/@ProGameplayDota2          No transcript available 
#     "UCvTcxoyItMUSlw8T2MajftA",   # BananaSlamJamma                                    OK
#     "UCUqLL4VcEy4mXcQL0O_H_bg",   # DotA Digest                                        OK
#     "UCzvDSoNrl2Uva7t_Z936VUw",   # Dota 2 Tips, Tricks and Guides                     too old 
#     "UCy0-ftAwxMHzZc74OhYT3PA",   # GameLeap Dota 2 Pro Guides                         too old 
#     "UCUqLL4VcEy4mXcQL0O_H_bg",   # DotA Digest                                    
# ]                                 

import os
import re
import time
import tempfile
from datetime import datetime, timedelta
from urllib.parse import urlparse, parse_qs
import streamlit as st
from googleapiclient.discovery import build
from youtube_transcript_api import (
    YouTubeTranscriptApi,
    TranscriptsDisabled,
    NoTranscriptFound,
    VideoUnavailable
)
from docx import Document
from docx.shared import Pt

# Set the page configuration with the custom favicon
# st.set_page_config(
#     page_title="Crazy Horse",       # Title that appears in the browser tab
#     page_icon="crazy-horse-icon.png",           # Path to your favicon image
# )
# st.image("crazy-horse.png", width=400)

col1, col2, col3 = st.columns([1, 2, 1])

# Place the image in the center column
with col2:
    st.image("crazy-horse.png", use_column_width=True)


# ----------------------------
# Configuration Variables
# ----------------------------
API_SERVICE_NAME = 'youtube'
API_VERSION = 'v3'

# ----------------------------
# Translations Dictionary
# ----------------------------
translations = {
    'en': {
        'title': 'YouTube Transcript Scraper',
        'description': 'This application allows you to fetch transcripts from a YouTube channel\'s videos or individual videos. Enter one or more channel URLs, video URLs, or channel handles (e.g., @ChannelName), specify the number of transcripts to fetch, and set a cutoff publish date.',
        'api_key': 'YouTube API Key',
        'api_key_help': 'Enter your YouTube Data API key.',
        'select_language': 'Select Language',
        'channel_url': 'YouTube Channel URL, Video URL, or Channel Handle (e.g., @ChannelName) (separated by commas or newlines)',
        'num_transcripts': 'Number of Transcripts to Fetch (0 for all, leave blank for 1)',
        'cutoff_date': 'Cutoff Publish Date',
        'fetch_button': 'Fetch Transcripts',
        'invalid_channel_url': 'Invalid YouTube channel URL.',
        'invalid_video_url': 'Invalid YouTube video URL.',
        'invalid_channel_id_format': 'Invalid channel ID format.',
        'unsupported_url_format': 'Unsupported YouTube channel URL format.',
        'error_fetching_channel_id': 'Error fetching channel ID: {}',
        'error_fetching_videos': 'Error fetching videos for channel: {}',
        'error_fetching_video_details': 'Error fetching video details: {}',
        'error_fetching_channel_id_from_video': 'Error fetching channel ID from video: {}',
        'manual_transcript_found': 'Manually created transcript found for video "{}" in language "{}".',
        'no_manual_transcript': 'No manually created transcript found for video "{}" in languages {}.',
        'auto_transcript_found': 'Auto-generated transcript found for video "{}" in language "{}".',
        'transcript_found': 'Transcript found for video "{}" in language "{}".',
        'no_transcript_available': 'No transcript available for video "{}" in languages {}.',
        'transcripts_disabled': 'Transcripts are disabled for video "{}".',
        'video_unavailable': 'Video "{}" is unavailable.',
        'unexpected_error': 'An unexpected error occurred while fetching transcript for video "{}": {}',
        'summary_prepared': 'Summary of all videos prepared.',
        'transcripts_to_process': 'Transcripts to process: {}',
        'transcript_fetched': 'Transcript fetched for video: "{}".',
        'skipping_video': 'Skipping video ID "{}" due to unavailable transcript.',
        'completed_processing': 'Completed processing. Total transcripts fetched: {}',
        'concatenated_docx_created': 'Concatenated transcript document created.',
        'download_results': 'Download Results',
        'download_summary': 'Download Summary',
        'download_docx': 'Download Transcripts (.docx)',
        'processing_channel_url': 'Processing channel input: {}',
        'found_videos': 'Found {} videos in channel \'{}\'',
        'no_videos_found': 'No videos found. Exiting.',
        'found_videos_after_cutoff': 'Found {} videos published after {}',
        'summary_saved': 'Summary of all videos saved to {}',
        'fetching_video_ids': 'Fetching video IDs...',
        'fetching_video_details': 'Fetching video details...',
        'fetching_transcript': 'Fetching transcript {}/{} for video ID: {}',
        'progress': 'Processed {}/{} transcripts.',
        'error_writing_summary': 'Error writing summary: {}',
        'error_creating_docx': 'Error creating concatenated .docx file: {}',
        'channel_name': 'Channel Name: {}',
        'no_video_details_found': 'No video details found. Exiting.',
        'multiple_channels_error': 'Multiple channels detected from the provided inputs. Please ensure all inputs belong to the same channel.',
        'no_valid_urls': 'No valid YouTube channel, video URLs, or channel handles provided.',
        'found_videos_specific': 'Found {} videos to process.',
        'no_transcripts_fetched': 'No transcripts were fetched.',
        'invalid_url': 'Invalid YouTube URL or handle: {}',
        'api_key_missing': 'Please enter your YouTube API Key to proceed.',
        'channel_url_help': 'Enter one or more YouTube channel URLs, video URLs, or channel handles (e.g., @ChannelName), separated by commas or newlines.',
        'api_key_instructions': """
**How to obtain a YouTube Data API Key:**

1. Go to the [Google Cloud Console](https://console.cloud.google.com/).
2. Sign in with your Google account.
3. Click on the **Select a project** dropdown at the top and choose **New Project**.
4. Enter a project name and click **Create**.
5. Once the project is created, ensure it's selected.
6. Navigate to the **APIs & Services** dashboard from the left sidebar.
7. Click on **Enable APIs and Services**.
8. Search for **YouTube Data API v3** and select it.
9. Click on **Enable**.
10. After enabling, go to the **Credentials** tab.
11. Click on **Create Credentials** > **API key**.
12. Your new API key will be displayed. Copy it and paste it into the app.
"""
    },
    'fi': {
        'title': 'YouTube tekstitystiedostojen hakija',
        'description': 'Tämä sovellus mahdollistaa tekstitystiedostojen hakemisen YouTube-kanavan videoista tai yksittäisistä videoista. Syötä yksi tai useampi kanavan URL-osoite, videon URL-osoite tai kanavan nimi (esim. @ChannelName), määritä haettavien tekstitystiedostojen määrä ja anna päivämäärä, jota vanhempia videoita ei haeta.',
        'api_key': 'YouTube API-avain',
        'api_key_help': 'Syötä YouTube Data API -avaimesi.',
        'select_language': 'Valitse Kieli',
        'channel_url': 'YouTube-kanavan URL-osoite, videon URL-osoite tai kanavan nimi (esim. @ChannelName) (erottele pilkuilla tai riveillä)',
        'num_transcripts': 'Haettavien tekstitystiedostojen määrä (0 -> haen kaikki videot, oletuksena 1 video)',
        'cutoff_date': 'Älä etsi tätä vanhempia videoita, anna päivämäärä',
        'fetch_button': 'Hae videoiden tekstitiedostot',
        'invalid_channel_url': 'Virheellinen YouTube-kanavan URL-osoite.',
        'invalid_video_url': 'Virheellinen YouTube-videon URL-osoite.',
        'invalid_channel_id_format': 'Virheellinen kanavan ID-muoto.',
        'unsupported_url_format': 'Tuettu YouTube-kanavan URL-osoiteformaatti.',
        'error_fetching_channel_id': 'Virhe kanavan ID:n hakemisessa: {}',
        'error_fetching_videos': 'Virhe videoiden hakemisessa kanavalle: {}',
        'error_fetching_video_details': 'Virhe videon tietojen hakemisessa: {}',
        'error_fetching_channel_id_from_video': 'Virhe kanavan ID:n hakemisessa videosta: {}',
        'manual_transcript_found': 'Manuaalisesti luotu tekstitystiedosto löytyi videolle "{}" kielillä "{}".',
        'no_manual_transcript': 'Manuaalista tekstitystiedostoa ei löytynyt videolle "{}" kielillä {}.',
        'auto_transcript_found': 'Automaattisesti luotu tekstitystiedosto löytyi videolle "{}" kielillä "{}".',
        'transcript_found': 'Tekstitystiedosto löytyi videolle "{}" kielillä "{}".',
        'no_transcript_available': 'Tekstitystiedostoa ei ole saatavilla videolle "{}" kielillä {}.',
        'transcripts_disabled': 'Tekstitystiedostot ovat pois käytöstä videolle "{}".',
        'video_unavailable': 'Video "{}" ei ole saatavilla.',
        'unexpected_error': 'Odottamaton virhe tekstitystiedostoa haettaessa videolle "{}": {}',
        'summary_prepared': 'Kaikkien videoiden yhteenveto valmisteltu.',
        'transcripts_to_process': 'Prosessoitavien tekstitystiedostojen määrä: {}',
        'transcript_fetched': 'Tekstitystiedosto haettu videolle: "{}".',
        'skipping_video': 'Ohitetaan video ID "{}" tekstitystiedoston puutteellisuuden vuoksi.',
        'completed_processing': 'Prosessointi valmis. Haetut tekstitystiedostot yhteensä: {}',
        'concatenated_docx_created': 'Yhdistetty tekstitystiedostosiakirja luotu.',
        'download_results': 'Lataa tulokset',
        'download_summary': 'Lataa Yhteenveto',
        'download_docx': 'Lataa Tekstitystiedostot (.docx)',
        'processing_channel_url': 'Prosessoidaan kanavan syöte: {}',
        'found_videos': 'Löytyi {} videota kanavasta \'{}\'',
        'no_videos_found': 'Videoita ei löytynyt. Poistutaan.',
        'found_videos_after_cutoff': 'Löytyi {} videota, jotka julkaistiin jälkeen {}',
        'summary_saved': 'Kaikkien videoiden yhteenveto tallennettu tiedostoon {}',
        'fetching_video_ids': 'Haetaan videoiden ID:tä...',
        'fetching_video_details': 'Haetaan videoiden tietoja...',
        'fetching_transcript': 'Haetaan tekstitystiedostoa {}/{} videolle ID: {}',
        'progress': 'Prosessoitu {}/{} tekstitystiedostoa.',
        'error_writing_summary': 'Virhe yhteenvetotiedoston kirjoittamisessa: {}',
        'error_creating_docx': 'Virhe yhdistetyn .docx-tiedoston luomisessa: {}',
        'channel_name': 'Kanavan Nimi: {}',
        'no_video_details_found': 'Videon tietoja ei löytynyt. Poistutaan.',
        'multiple_channels_error': 'Annetuista syötteistä löytyi useita kanavia. Varmista, että kaikki syötteet kuuluvat samaan kanavaan.',
        'no_valid_urls': 'Yhtään kelvollista YouTube-kanavan URL-osoitetta, videon URL-osoitetta tai kanavan nimeä ei annettu.',
        'found_videos_specific': 'Löytyi {} videota prosessoitavaksi.',
        'no_transcripts_fetched': 'Tekstitystiedostoja ei haettu.',
        'invalid_url': 'Virheellinen YouTube-URL tai nimi: {}',
        'api_key_missing': 'Please enter your YouTube API Key to proceed.',
        'channel_url_help': 'Syötä yksi tai useampi YouTube-kanavan URL-osoite, videon URL-osoite tai kanavan nimi (esim. @ChannelName), eroteltuna pilkuilla tai riveillä.',
        'api_key_instructions': """
**Kuinka hankkia YouTube Data API -avain:**

1. Mene [Google Cloud Consoleen](https://console.cloud.google.com/).
2. Kirjaudu sisään Google-tililläsi.
3. Napsauta ylhäällä olevaa **Valitse projekti** -pudotusvalikkoa ja valitse **Uusi projekti**.
4. Anna projektille nimi ja napsauta **Luo**.
5. Kun projekti on luotu, varmista että se on valittuna.
6. Siirry vasemmasta sivupalkista **APIs & Services** -hallintapaneeliin.
7. Napsauta **Ota API:t ja palvelut käyttöön**.
8. Etsi **YouTube Data API v3** ja valitse se.
9. Napsauta **Ota käyttöön**.
10. Käyttöönoton jälkeen siirry **Credentials**-välilehdelle.
11. Napsauta **Luo tunnistetiedot** > **API-avain**.
12. Uusi API-avaimesi näytetään. Kopioi se ja liitä se sovellukseen.
"""
    }
}

# ----------------------------
# Default selected language
# ----------------------------
selected_language = 'en'

def set_language():
    """
    Sets the selected language based on user input.
    """
    global selected_language
    lang = st.sidebar.selectbox(
        translations['en']['select_language'],
        options=['English', 'Finnish'],
        index=0,
        format_func=lambda x: x
    )
    selected_language = 'en' if lang == 'English' else 'fi'

def sanitize_filename(name):
    """
    Removes or replaces characters that are invalid in filenames.
    """
    return re.sub(r'[\\/:*?"<>|]', '_', name).strip()

def get_youtube_client(api_key):
    """
    Initializes and returns the YouTube API client.
    
    Args:
        api_key (str): The YouTube Data API key.
    
    Returns:
        googleapiclient.discovery.Resource: The YouTube API client.
    """
    try:
        youtube_client = build(API_SERVICE_NAME, API_VERSION, developerKey=api_key)
        st.sidebar.write("YouTube API client initialized successfully.")  # Debugging statement
        return youtube_client
    except Exception as e:
        st.sidebar.error(translations[selected_language]['error_fetching_channel_id'].format(e))
        return None

def get_channel_id_from_identifier(youtube, identifier, identifier_type):
    """
    Retrieves the channel ID based on identifier type.
    
    Args:
        youtube: YouTube API client.
        identifier (str): The username, custom name, or handle.
        identifier_type (str): 'user', 'custom', or 'handle'

    Returns:
        str or None: The channel ID if found, else None.
    """
    try:
        if identifier_type == 'user':
            response = youtube.channels().list(
                part='id',
                forUsername=identifier
            ).execute()
            if response['items']:
                return response['items'][0]['id']
        elif identifier_type == 'custom':
            response = youtube.search().list(
                part='snippet',
                q=identifier,
                type='channel',
                maxResults=1
            ).execute()
            if response['items']:
                return response['items'][0]['snippet']['channelId']
        elif identifier_type == 'handle':
            # Remove '@' from handle
            handle = identifier.lstrip('@')
            response = youtube.search().list(
                part='snippet',
                q=handle,
                type='channel',
                maxResults=1
            ).execute()
            if response['items']:
                return response['items'][0]['snippet']['channelId']
    except Exception as e:
        st.sidebar.error(translations[selected_language]['error_fetching_channel_id'].format(e))
    return None

def extract_channel_id(youtube, input_str):
    """
    Extracts the channel ID from a YouTube channel URL or handle.

    Args:
        youtube: YouTube API client.
        input_str (str): The YouTube channel URL or handle.

    Returns:
        str or None: The channel ID if found, else None.
    """
    input_str = input_str.strip()
    if input_str.startswith('@'):
        # It's a handle, construct the full URL
        url = f'https://www.youtube.com/{input_str}'
        st.sidebar.write(f"Detected handle: {input_str}. Constructed URL: {url}")
    else:
        url = input_str

    parsed_url = urlparse(url)
    path = parsed_url.path
    # Remove trailing slash
    if path.endswith('/'):
        path = path[:-1]
    parts = path.split('/')

    if len(parts) < 2:
        st.sidebar.error(translations[selected_language]['invalid_channel_url'])
        return None

    identifier_type = parts[1]
    identifier = parts[2] if len(parts) > 2 else ''

    if identifier_type == 'channel':
        # URL format: /channel/UCxxxxx
        channel_id = identifier
        if re.match(r'^UC[a-zA-Z0-9_-]{22}$', channel_id):
            return channel_id
        else:
            st.sidebar.error(translations[selected_language]['invalid_channel_id_format'])
            return None
    elif identifier_type == 'user':
        # URL format: /user/username
        return get_channel_id_from_identifier(youtube, identifier, 'user')
    elif identifier_type == 'c':
        # URL format: /c/customname
        return get_channel_id_from_identifier(youtube, identifier, 'custom')
    elif identifier_type.startswith('@'):
        # URL format: /@handle
        return get_channel_id_from_identifier(youtube, identifier_type, 'handle')
    else:
        st.sidebar.error(translations[selected_language]['unsupported_url_format'])
        return None

def extract_video_id(url):
    """
    Extracts the video ID from a YouTube video URL.

    Args:
        url (str): The YouTube video URL.

    Returns:
        str or None: The video ID if found, else None.
    """
    parsed_url = urlparse(url.strip())
    if parsed_url.hostname in ['www.youtube.com', 'youtube.com']:
        query = parse_qs(parsed_url.query)
        if 'v' in query:
            return query['v'][0]
    elif parsed_url.hostname in ['youtu.be']:
        return parsed_url.path.lstrip('/')
    st.sidebar.error(translations[selected_language]['invalid_video_url'].format(url))
    return None

def get_channel_id_from_video(youtube, video_id):
    """
    Retrieves the channel ID from a video ID.

    Args:
        youtube: YouTube API client.
        video_id (str): The YouTube video ID.

    Returns:
        str or None: The channel ID if found, else None.
    """
    try:
        response = youtube.videos().list(
            part='snippet',
            id=video_id
        ).execute()
        items = response.get('items', [])
        if not items:
            st.sidebar.error(translations[selected_language]['video_unavailable'].format(video_id))
            return None
        channel_id = items[0]['snippet']['channelId']
        return channel_id
    except Exception as e:
        st.sidebar.error(translations[selected_language]['error_fetching_channel_id_from_video'].format(e))
        return None

def get_all_video_ids(youtube, channel_id):
    """
    Retrieves all video IDs from a given channel ID.

    Args:
        youtube: The YouTube API client.
        channel_id (str): The YouTube channel ID.

    Returns:
        list: A list of video IDs.
    """
    video_ids = []
    try:
        # Get uploads playlist ID
        response = youtube.channels().list(
            part='contentDetails',
            id=channel_id
        ).execute()
        uploads_playlist_id = response['items'][0]['contentDetails']['relatedPlaylists']['uploads']

        next_page_token = None
        while True:
            playlist_response = youtube.playlistItems().list(
                part='contentDetails',
                playlistId=uploads_playlist_id,
                maxResults=50,
                pageToken=next_page_token
            ).execute()

            for item in playlist_response['items']:
                video_ids.append(item['contentDetails']['videoId'])

            next_page_token = playlist_response.get('nextPageToken')
            if not next_page_token:
                break
            # To respect API rate limits
            time.sleep(0.1)
    except Exception as e:
        st.sidebar.error(translations[selected_language]['error_fetching_videos'].format(e))
    return video_ids

def fetch_transcript(video_id):
    """
    Fetches the transcript for a given YouTube video ID.
    Prioritizes transcripts in fi, then sv, then en. If none are available, selects any available transcript.

    Args:
        video_id (str): The YouTube video ID.

    Returns:
        list or None: The transcript as a list of dictionaries if available, else None.
    """
    preferred_languages = ['fi', 'sv', 'en']
    try:
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)

        # Attempt to find a manually created transcript in preferred languages
        for lang in preferred_languages:
            try:
                transcript = transcript_list.find_manually_created_transcript([lang])
                st.info(translations[selected_language]['manual_transcript_found'].format(video_id, lang))
                return transcript.fetch()
            except NoTranscriptFound:
                continue  # Try next preferred language

        # If no manually created transcripts found in preferred languages, try auto-generated
        for lang in preferred_languages:
            try:
                transcript = transcript_list.find_generated_transcript([lang])
                st.info(translations[selected_language]['auto_transcript_found'].format(video_id, lang))
                return transcript.fetch()
            except NoTranscriptFound:
                continue  # Try next preferred language

        # If no preferred language transcripts are found, select any available transcript
        try:
            transcript = transcript_list.find_transcript(transcript_list._list_transcripts)
            language = transcript.language
            transcript_type = 'auto-generated' if transcript.is_generated else 'manually created'
            st.info(translations[selected_language]['transcript_found'].format(video_id, language))
            return transcript.fetch()
        except NoTranscriptFound:
            st.warning(translations[selected_language]['no_transcript_available'].format(video_id, preferred_languages))
            return None

    except TranscriptsDisabled:
        st.warning(translations[selected_language]['transcripts_disabled'].format(video_id))
        return None
    except VideoUnavailable:
        st.warning(translations[selected_language]['video_unavailable'].format(video_id))
        return None
    except Exception as e:
        st.error(translations[selected_language]['unexpected_error'].format(video_id, e))
        return None

def transcript_to_text(transcript):
    """
    Converts the transcript list to plain text.

    Args:
        transcript (list): The transcript as a list of dictionaries.

    Returns:
        str: The transcript in plain text format.
    """
    return "\n".join([entry['text'] for entry in transcript])

def get_video_details(youtube, video_ids):
    """
    Retrieves the title and publish datetime of videos.

    Args:
        youtube: The YouTube API client.
        video_ids (list): List of YouTube video IDs.

    Returns:
        list of dict: Each dict contains 'video_id', 'title', and 'publish_datetime'.
    """
    details = []
    try:
        for i in range(0, len(video_ids), 50):
            batch_ids = video_ids[i:i+50]
            response = youtube.videos().list(
                part='snippet',
                id=','.join(batch_ids)
            ).execute()
            for item in response.get('items', []):
                video_id = item['id']
                title = item['snippet']['title']
                publish_time = item['snippet']['publishedAt']
                # Convert publish_time to desired format
                publish_datetime = datetime.strptime(publish_time, "%Y-%m-%dT%H:%M:%SZ")
                formatted_datetime = publish_datetime.strftime("%Y-%m-%d-%H-%M")
                details.append({
                    'video_id': video_id,
                    'title': title,
                    'publish_datetime': formatted_datetime
                })
    except Exception as e:
        st.sidebar.error(translations[selected_language]['error_fetching_video_details'].format(e))
    return details

def create_concatenated_docx(transcripts, document):
    """
    Creates a concatenated .docx file with each transcript preceded by a heading.

    Args:
        transcripts (list of dict): Each dict contains 'publish_datetime', 'title', and 'text'.
        document (Document): The python-docx Document object to append content to.
    """
    for transcript in transcripts:
        publish_datetime = transcript['publish_datetime']
        title = transcript['title']
        text = transcript['text']

        # Add heading
        heading = f"{publish_datetime} - {title}"
        document.add_heading(heading, level=1)

        # Add transcript text
        paragraph = document.add_paragraph(text)
        paragraph.style.font.size = Pt(12)

def main(urls, num_transcripts, cutoff_date, youtube_client):
    """
    Main function to process transcripts from YouTube channels or individual videos.

    Args:
        urls (list): List of YouTube channel URLs, video URLs, or channel handles.
        num_transcripts (int or None): Number of transcripts to fetch. If 0, fetch all.
        cutoff_date (datetime): The cutoff publish date.
        youtube_client (googleapiclient.discovery.Resource): The YouTube API client.

    Returns:
        tuple: (summary_bytes, concatenated_docx_bytes)
    """
    st.header(translations[selected_language]['title'])

    st.write(translations[selected_language]['processing_channel_url'].format(', '.join(urls)))

    channel_ids = set()
    video_ids_specific = []

    # Process each input
    for input_str in urls:
        input_str = input_str.strip()
        if not input_str:
            continue

        if input_str.startswith('@'):
            # It's a handle, construct the full URL
            url = f'https://www.youtube.com/{input_str}'
            st.write(f"Detected handle: {input_str}. Constructed URL: {url}")
        else:
            url = input_str

        parsed_url = urlparse(url)
        if 'youtube.com' not in parsed_url.netloc and 'youtu.be' not in parsed_url.netloc:
            st.error(translations[selected_language]['invalid_url'].format(input_str))
            return None, None

        if '/watch' in parsed_url.path or parsed_url.hostname in ['youtu.be']:
            # It's a video URL
            video_id = extract_video_id(url)
            if not video_id:
                st.error(translations[selected_language]['invalid_video_url'].format(url))
                return None, None
            channel_id = get_channel_id_from_video(youtube_client, video_id)
            if channel_id:
                channel_ids.add(channel_id)
                video_ids_specific.append(video_id)
        else:
            # Assume it's a channel URL or handle
            channel_id = extract_channel_id(youtube_client, input_str)
            if channel_id:
                channel_ids.add(channel_id)

    # Validate channel IDs
    if len(channel_ids) > 1:
        st.error(translations[selected_language]['multiple_channels_error'])
        return None, None
    elif len(channel_ids) == 0:
        st.error(translations[selected_language]['no_valid_urls'])
        return None, None

    channel_id = channel_ids.pop()

    # Get channel details for naming
    try:
        channel_response = youtube_client.channels().list(
            part='snippet',
            id=channel_id
        ).execute()
        channel_name = channel_response['items'][0]['snippet']['title']
        sanitized_channel_name = sanitize_filename(channel_name)
    except Exception as e:
        st.error(translations[selected_language]['error_fetching_channel_id'].format(e))
        channel_name = channel_id  # Fallback to channel ID
        sanitized_channel_name = sanitize_filename(channel_id)

    st.success(translations[selected_language]['channel_name'].format(channel_name))

    # Determine mode: channel or specific videos
    if video_ids_specific:
        # User provided video URLs or handles that map to specific videos
        video_ids = video_ids_specific
        st.write(translations[selected_language]['found_videos_specific'].format(len(video_ids)))
    else:
        # User provided channel URLs or handles
        with st.spinner(translations[selected_language]['fetching_video_ids']):
            video_ids = get_all_video_ids(youtube_client, channel_id)
        st.write(translations[selected_language]['found_videos'].format(len(video_ids), channel_name))

        if not video_ids:
            st.warning(translations[selected_language]['no_videos_found'])
            return None, None

    # Initialize variables to hold summary and transcripts
    summary_text = ""
    concatenated_transcripts = []

    # If channel mode, get video details
    if not video_ids_specific:
        with st.spinner(translations[selected_language]['fetching_video_details']):
            video_details = get_video_details(youtube_client, video_ids)
        if not video_details:
            st.warning(translations[selected_language]['no_video_details_found'])
            return None, None

        # Sort videos from latest to oldest
        video_details_sorted = sorted(video_details, key=lambda x: x['publish_datetime'], reverse=True)

        # Filter videos based on publish date (only transcripts for videos published after cutoff_date)
        filtered_videos = [
            video for video in video_details_sorted
            if datetime.strptime(video['publish_datetime'], "%Y-%m-%d-%H-%M") > cutoff_date
        ]

        st.write(translations[selected_language]['found_videos_after_cutoff'].format(len(filtered_videos), cutoff_date.strftime('%Y-%m-%d')))

        # Write summary.txt with all videos
        for video in video_details_sorted:
            line = f"{video['publish_datetime']} - {video['title']}\n"
            summary_text += line

        # Determine transcripts to process based on num_transcripts
        if num_transcripts == 0:
            transcripts_to_process = filtered_videos
        else:
            transcripts_to_process = filtered_videos[:num_transcripts]

        st.write(translations[selected_language]['transcripts_to_process'].format(len(transcripts_to_process)))

        # Initialize file counter
        file_counter = 0

        # Progress bar
        progress_bar = st.progress(0)
        progress_text = st.empty()

        for idx, video in enumerate(transcripts_to_process, 1):
            video_id = video['video_id']
            publish_datetime = video['publish_datetime']
            video_title = video['title']
            st.write(translations[selected_language]['fetching_transcript'].format(idx, len(transcripts_to_process), video_id))
            transcript = fetch_transcript(video_id)
            if transcript:
                text = transcript_to_text(transcript)
                concatenated_transcripts.append({
                    'publish_datetime': publish_datetime,
                    'title': video_title,
                    'text': text
                })
                file_counter += 1
                st.success(translations[selected_language]['transcript_fetched'].format(video_title))
            else:
                st.warning(translations[selected_language]['skipping_video'].format(video_id))

            # Update progress bar
            progress = idx / len(transcripts_to_process)
            progress_bar.progress(progress)
            progress_text.text(translations[selected_language]['progress'].format(idx, len(transcripts_to_process)))

            # To avoid hitting rate limits
            time.sleep(1)

        st.success(translations[selected_language]['completed_processing'].format(file_counter))

    else:
        # Video mode
        with st.spinner(translations[selected_language]['fetching_video_details']):
            video_details = get_video_details(youtube_client, video_ids)
        if not video_details:
            st.warning(translations[selected_language]['no_video_details_found'])
            return None, None

        # Sort videos from latest to oldest
        video_details_sorted = sorted(video_details, key=lambda x: x['publish_datetime'], reverse=True)

        # Filter videos based on publish date (only transcripts for videos published after cutoff_date)
        filtered_videos = [
            video for video in video_details_sorted
            if datetime.strptime(video['publish_datetime'], "%Y-%m-%d-%H-%M") > cutoff_date
        ]

        st.write(translations[selected_language]['found_videos_after_cutoff'].format(len(filtered_videos), cutoff_date.strftime('%Y-%m-%d')))

        # Write summary.txt with all videos
        for video in video_details_sorted:
            line = f"{video['publish_datetime']} - {video['title']}\n"
            summary_text += line

        # Determine transcripts to process based on num_transcripts
        if num_transcripts == 0:
            transcripts_to_process = filtered_videos
        else:
            transcripts_to_process = filtered_videos[:num_transcripts]

        st.write(translations[selected_language]['transcripts_to_process'].format(len(transcripts_to_process)))

        # Initialize file counter
        file_counter = 0

        # Progress bar
        progress_bar = st.progress(0)
        progress_text = st.empty()

        for idx, video in enumerate(transcripts_to_process, 1):
            video_id = video['video_id']
            publish_datetime = video['publish_datetime']
            video_title = video['title']
            st.write(translations[selected_language]['fetching_transcript'].format(idx, len(transcripts_to_process), video_id))
            transcript = fetch_transcript(video_id)
            if transcript:
                text = transcript_to_text(transcript)
                concatenated_transcripts.append({
                    'publish_datetime': publish_datetime,
                    'title': video_title,
                    'text': text
                })
                file_counter += 1
                st.success(translations[selected_language]['transcript_fetched'].format(video_title))
            else:
                st.warning(translations[selected_language]['skipping_video'].format(video_id))

            # Update progress bar
            progress = idx / len(transcripts_to_process)
            progress_bar.progress(progress)
            progress_text.text(translations[selected_language]['progress'].format(idx, len(transcripts_to_process)))

            # To avoid hitting rate limits
            time.sleep(1)

        st.success(translations[selected_language]['completed_processing'].format(file_counter))

    # Use a temporary directory to store transcripts and summary
    with tempfile.TemporaryDirectory() as temp_dir:
        summary_file_path = os.path.join(temp_dir, 'summary.txt')
        try:
            with open(summary_file_path, 'w', encoding='utf-8') as summary_file:
                summary_file.write(summary_text)
            st.success(translations[selected_language]['summary_prepared'])
        except Exception as e:
            st.error(translations[selected_language]['error_writing_summary'].format(e))

        # Create concatenated docx file
        concatenated_docx_bytes = None  # Initialize
        if concatenated_transcripts:
            document = Document()
            document.add_heading(f"Transcripts for {channel_name}", 0)
            create_concatenated_docx(concatenated_transcripts, document)

            concatenated_docx_path = os.path.join(temp_dir, 'transcripts.docx')
            try:
                document.save(concatenated_docx_path)
                st.success(translations[selected_language]['concatenated_docx_created'])
                # Read the docx file into bytes
                with open(concatenated_docx_path, 'rb') as docx_file:
                    concatenated_docx_bytes = docx_file.read()
            except Exception as e:
                st.error(translations[selected_language]['error_creating_docx'].format(e))
        else:
            st.warning(translations[selected_language]['no_transcripts_fetched'])

        # Read the summary file into bytes
        summary_bytes = None
        if os.path.exists(summary_file_path):
            with open(summary_file_path, 'rb') as summary_file:
                summary_bytes = summary_file.read()

        return summary_bytes, concatenated_docx_bytes

def main_entry():
    # Initialize session state variables if they don't exist
    if 'summary_bytes' not in st.session_state:
        st.session_state.summary_bytes = None
    if 'concatenated_docx_bytes' not in st.session_state:
        st.session_state.concatenated_docx_bytes = None
    if 'show_api_key_instructions' not in st.session_state:
        st.session_state.show_api_key_instructions = False

    # Sidebar inputs
    st.sidebar.header("Settings")

    # YouTube API Key Input with Info button
    col1, col2 = st.sidebar.columns([4, 1])
    with col1:
        api_key = st.sidebar.text_input(
            translations['en']['api_key'],  # Always display in English
            type='password',
            help=translations[selected_language]['api_key_help']
        )
    with col2:
        if st.sidebar.button("ℹ️"):
            st.session_state.show_api_key_instructions = True

    # Display instructions if flag is True
    if st.session_state.show_api_key_instructions:
        st.sidebar.markdown(translations[selected_language]['api_key_instructions'])
        if st.sidebar.button("Close Instructions"):
            st.session_state.show_api_key_instructions = False

    # Language Selection
    set_language()

    # If API key is not provided, show an error and stop further execution
    if not api_key:
        st.error(translations[selected_language]['api_key_missing'])
        st.stop()

    # Initialize YouTube API client
    youtube_client = get_youtube_client(api_key)
    if not youtube_client:
        st.stop()

    st.title(translations[selected_language]['title'])

    st.write(translations[selected_language]['description'])

    # User inputs
    urls_input = st.text_area(
        translations[selected_language]['channel_url'],
        height=150,
        value="@SabineHossenfelder", 
        help=translations[selected_language].get('channel_url_help', '')
    )
    urls = re.split(r'[,\n]+', urls_input) if urls_input else []
    urls = [url.strip() for url in urls if url.strip()]

    num_transcripts = st.number_input(
        translations[selected_language]['num_transcripts'],
        min_value=0,
        step=1,
        value=1,  # Set default value to 1
        format="%d"
    )
    # Handle 'leave blank for 1' by setting to 1 if left blank
    num_transcripts = int(num_transcripts) if num_transcripts else 1
    cutoff_date = st.date_input(
        translations[selected_language]['cutoff_date'],
        value=(datetime.now() - timedelta(days=2*365)).date()
    )

    if st.button(translations[selected_language]['fetch_button']):
        if not urls:
            st.error(translations[selected_language]['no_valid_urls'])
        else:
            # Convert cutoff_date from datetime.date to datetime.datetime
            cutoff_datetime = datetime.combine(cutoff_date, datetime.min.time())
            # Clear previous session state data
            st.session_state.summary_bytes = None
            st.session_state.concatenated_docx_bytes = None
            # Run main processing
            summary_bytes, concatenated_docx_bytes = main(urls, num_transcripts, cutoff_datetime, youtube_client)
            # Save the summary and docx bytes to session_state
            st.session_state.summary_bytes = summary_bytes
            st.session_state.concatenated_docx_bytes = concatenated_docx_bytes

    # Display download buttons if data is available in session_state
    if st.session_state.summary_bytes or st.session_state.concatenated_docx_bytes:
        st.header(translations[selected_language]['download_results'])

        # Summary file download
        if st.session_state.summary_bytes:
            st.download_button(
                label=translations[selected_language]['download_summary'],
                data=st.session_state.summary_bytes,
                file_name='summary.txt',
                mime='text/plain'
            )

        # Concatenated transcripts download
        if st.session_state.concatenated_docx_bytes:
            st.download_button(
                label=translations[selected_language]['download_docx'],
                data=st.session_state.concatenated_docx_bytes,
                file_name='transcripts.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

if __name__ == "__main__":
    main_entry()
